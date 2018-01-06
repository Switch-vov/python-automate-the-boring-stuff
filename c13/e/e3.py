import docx

doc = docx.Document()
doc.add_paragraph('Hello World!')
paragraphs = doc.paragraphs[0]
paragraphs.style = 'Title'
doc.save('hello world.docx')
